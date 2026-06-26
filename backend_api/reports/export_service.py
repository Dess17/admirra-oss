"""
Сервис экспорта отчётов: PNG, DOCX.
Использует данные из pdf_service и генерирует альтернативные форматы.
"""
import logging
import os
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional
import uuid as uuid_module

from sqlalchemy.orm import Session

from backend_api.reports.pdf_service import generate_report_pdf
from backend_api.stats_service import StatsService
from core import models

logger = logging.getLogger(__name__)
VAT_RATE = 1.22

# Временное хранилище файлов для ссылок: token -> (filepath, expires_at)
_report_file_cache: dict[str, tuple[str, float]] = {}

# Хранилище данных для страницы просмотра отчёта: token -> (data_dict, expires_at)
_report_view_cache: dict[str, tuple[dict, float]] = {}


def _is_avito_platform(value) -> bool:
    return str(value or "").strip().lower() in {"avito", "avito_ads"}


def _campaign_platform(campaign: dict) -> str:
    platform = campaign.get("platform") or campaign.get("channel")
    if platform:
        return str(platform)
    name = str(campaign.get("name") or campaign.get("campaign_name") or "").lower()
    if name.startswith("[avito]") or name.startswith("[авито]"):
        return "avito"
    return ""


def _with_channel_vat(value, platform=None) -> float:
    raw = float(value or 0)
    return raw if _is_avito_platform(platform) else raw * VAT_RATE


def _with_cost_breakdown_vat(value, cost_by_platform: dict | None, platform=None) -> float:
    if isinstance(cost_by_platform, dict):
        return (
            float(cost_by_platform.get("yandex") or 0) * VAT_RATE
            + float(cost_by_platform.get("vk") or 0) * VAT_RATE
            + float(cost_by_platform.get("avito") or 0)
        )
    return _with_channel_vat(value, platform)


def _summary_platform(campaigns: list) -> str:
    if campaigns and all(_is_avito_platform(_campaign_platform(c)) for c in campaigns):
        return "avito"
    return ""


def _get_report_data(db, user_id, client_id, start_date, end_date, comment):
    """Общие данные для отчёта (из pdf_service)."""
    effective_client_ids = StatsService.get_effective_client_ids(db, user_id, client_id)
    if not effective_client_ids:
        raise ValueError("Нет доступа к данным")
    try:
        d_end = datetime.strptime(end_date, "%Y-%m-%d").date()
        d_start = datetime.strptime(start_date, "%Y-%m-%d").date()
    except ValueError:
        raise ValueError("Неверный формат дат. Используйте YYYY-MM-DD.")
    summary = StatsService.aggregate_summary(
        db, effective_client_ids, d_start, d_end, "all", None, None
    )
    campaigns = StatsService.get_campaign_stats(
        db, effective_client_ids, d_start, d_end, "all", None, None
    )
    top_campaigns = sorted(
        [c for c in campaigns if c.get("conversions", 0) > 0],
        key=lambda x: x.get("conversions", 0),
        reverse=True,
    )[:10]
    client_name = None
    if client_id and len(effective_client_ids) == 1:
        client = db.query(models.Client).filter_by(id=client_id).first()
        if client:
            client_name = client.name
    ai_comment = (comment or "").strip() if comment else ""
    return summary, top_campaigns, client_name, ai_comment, start_date, end_date


def generate_report_png(
    db: Session,
    user_id: uuid_module.UUID,
    client_id: Optional[uuid_module.UUID],
    start_date: str,
    end_date: str,
    comment: Optional[str] = None,
) -> bytes:
    """Генерирует PNG отчёта (первая страница PDF)."""
    pdf_bytes = generate_report_pdf(
        db, user_id, client_id, start_date, end_date, comment
    )
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[0]
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        png_bytes = pix.tobytes("png")
        doc.close()
        return png_bytes
    except ImportError:
        logger.error("PyMuPDF (fitz) not installed")
        raise ImportError("Установите pymupdf: pip install pymupdf")


def generate_report_docx(
    db: Session,
    user_id: uuid_module.UUID,
    client_id: Optional[uuid_module.UUID],
    start_date: str,
    end_date: str,
    comment: Optional[str] = None,
) -> bytes:
    """Генерирует DOCX отчёт."""
    summary, top_campaigns, client_name, ai_comment, sd, ed = _get_report_data(
        db, user_id, client_id, start_date, end_date, comment
    )
    summary_platform = _summary_platform(top_campaigns)
    summary_expenses = _with_cost_breakdown_vat(summary.get("expenses", 0), summary.get("cost_by_platform"), summary_platform)
    summary_cpc = summary_expenses / float(summary.get("clicks") or 0) if summary.get("clicks") else _with_channel_vat(summary.get("cpc", 0), summary_platform)
    summary_cpa = summary_expenses / float(summary.get("leads") or 0) if summary.get("leads") else _with_channel_vat(summary.get("cpa", 0), summary_platform)
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading("Отчёт по рекламным кампаниям", 0)
        meta = f"Период: {sd} — {ed}"
        if client_name:
            meta += f" | Проект: {client_name}"
        doc.add_paragraph(meta).alignment = WD_ALIGN_PARAGRAPH.LEFT

        if ai_comment:
            doc.add_heading("Комментарий ИИ к отчёту", level=1)
            doc.add_paragraph(ai_comment)

        doc.add_heading("Ключевые показатели", level=1)
        kpi_text = (
            f"Расходы: {int(summary_expenses)} ₽ | "
            f"Показы: {int(summary.get('impressions', 0))} | "
            f"Клики: {int(summary.get('clicks', 0))} | "
            f"Лиды: {int(summary.get('leads', 0))} | "
            f"CPC: {summary_cpc:.2f} ₽ | "
            f"CPA: {summary_cpa:.2f} ₽"
        )
        doc.add_paragraph(kpi_text)

        if top_campaigns:
            doc.add_heading("Топ кампаний по конверсиям", level=1)
            table = doc.add_table(rows=1 + len(top_campaigns), cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Кампания"
            hdr[1].text = "Лиды"
            hdr[2].text = "Расход"
            hdr[3].text = "CPA"
            for i, c in enumerate(top_campaigns, 1):
                c_platform = _campaign_platform(c)
                row = table.rows[i].cells
                row[0].text = c.get("name", c.get("campaign_name", "—"))
                row[1].text = str(c.get("conversions", 0))
                row[2].text = f"{_with_channel_vat(c.get('cost', 0), c_platform):,.0f} ₽"
                row[3].text = f"{_with_channel_vat(c.get('cpa', 0), c_platform):.2f} ₽" if c.get("conversions") else "—"

        doc.add_paragraph(f"Сформировано: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        from io import BytesIO
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("Установите python-docx: pip install python-docx")


def save_report_for_link(
    format_type: str,
    file_bytes: bytes,
    start_date: str,
    end_date: str,
    ttl_seconds: int = 86400,
) -> str:
    """Сохраняет файл во временное хранилище, возвращает токен."""
    import time
    token = str(uuid.uuid4())
    ext = {"pdf": ".pdf", "png": ".png", "docx": ".docx"}.get(format_type, ".pdf")
    fd, path = tempfile.mkstemp(suffix=ext, prefix="report_")
    try:
        os.write(fd, file_bytes)
        os.close(fd)
        fd = None
        expires = time.time() + ttl_seconds
        _report_file_cache[token] = (path, expires)
        return token
    finally:
        if fd is not None:
            os.close(fd)


def save_report_view_data(
    summary: dict,
    top_campaigns: list,
    client_name: Optional[str],
    ai_comment: str,
    start_date: str,
    end_date: str,
    ttl_seconds: int = 86400,
) -> str:
    """Сохраняет данные отчёта для страницы просмотра, возвращает токен."""
    import time
    token = str(uuid.uuid4())
    data = {
        "summary": summary,
        "top_campaigns": top_campaigns,
        "client_name": client_name or "",
        "ai_comment": ai_comment,
        "start_date": start_date,
        "end_date": end_date,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    expires = time.time() + ttl_seconds
    _report_view_cache[token] = (data, expires)
    return token


def get_report_view_data(token: str) -> Optional[dict]:
    """Возвращает данные отчёта по токену или None если истёк."""
    import time
    if token not in _report_view_cache:
        return None
    data, expires = _report_view_cache[token]
    if time.time() > expires:
        del _report_view_cache[token]
        return None
    return data




def get_report_file_by_token(token: str) -> tuple[Optional[bytes], Optional[str], Optional[str]]:
    """Возвращает (bytes, media_type, filename) или (None, None, None) если не найден/истёк."""
    import time
    if token not in _report_file_cache:
        return None, None, None
    path, expires = _report_file_cache[token]
    if time.time() > expires:
        del _report_file_cache[token]
        try:
            os.unlink(path)
        except OSError:
            pass
        return None, None, None
    try:
        with open(path, "rb") as f:
            data = f.read()
        media_types = {
            ".pdf": "application/pdf",
            ".png": "image/png",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        ext = Path(path).suffix.lower()
        mt = media_types.get(ext, "application/octet-stream")
        fn = f"report{ext}"
        return data, mt, fn
    except OSError:
        return None, None, None
